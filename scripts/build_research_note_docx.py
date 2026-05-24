#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "FX検証研究ノート_2015-2024.docx"


def set_run_font(run, size_pt: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before: float = 0, after: float = 8, line: float = 1.15) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, 3, 1.15)
    r = p.add_run(title)
    set_run_font(r, 26, False, "000000")
    p2 = doc.add_paragraph()
    set_paragraph_spacing(p2, 0, 12, 1.15)
    r2 = p2.add_run(subtitle)
    set_run_font(r2, 11, False, "555555")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, 8, 1.15)
    r = p.add_run(text)
    set_run_font(r, 11, False, "000000")


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, 0, 4, 1.15)
    r = p.add_run(text)
    set_run_font(r, 11, False, "000000")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.style = "Table Grid"
    for idx, h in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_width(cell, widths[idx])
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, 0, 0, 1.15)
        r = p.add_run(h)
        set_run_font(r, 9.5, True, "000000")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, 0, 0, 1.15)
            r = p.add_run(str(value))
            set_run_font(r, 9.5, False, "000000")
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, 0, 8, 1.15)


def fmt_pct(v: float) -> str:
    return f"{v:.2f}%"


def fmt_r(v: float) -> str:
    return f"{v:+.2f}R"


def make_doc() -> None:
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, before, after, color in [
        ("Heading 1", 20, 20, 6, "000000"),
        ("Heading 2", 16, 18, 6, "000000"),
        ("Heading 3", 14, 16, 4, "434343"),
    ]:
        st = styles[name]
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        st.font.size = Pt(size)
        st.font.bold = False
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.15

    add_title(
        doc,
        "FX検証研究ノート（2015-2024）",
        "TrendBreakV1 / Sai近似 / V字フィボ / 急落後V字回復の検証メモ  |  作成日: 2026-05-23",
    )

    add_heading(doc, "1. このノートの目的", 1)
    add_body(
        doc,
        "ここまで行ったFX手法検証を、後から見返して再現できる研究ノートとして整理したものです。"
        "主な対象はTrendBreakV1、Sai手法の時間足比較、騙し回避、同方向追加ポジション、"
        "エリオット風5波、V字フィボ、急落後V字回復です。",
    )
    add_bullet(doc, "検証期間は主に2015年1月1日から2024年12月31日。")
    add_bullet(doc, "対象はXAUUSD、USDJPY、EURJPY、GBPJPY、CHFJPY、AUDJPY、SILVERなど。")
    add_bullet(doc, "結果はバックテスト上の研究値であり、将来の利益を保証するものではありません。")

    add_heading(doc, "2. 用語メモ", 1)
    add_table(
        doc,
        ["用語", "意味", "売買での使い方"],
        [
            ["R", "1回の許容損失を1単位にした成績表示。1%リスクなら+10Rは約+10%。", "資金額に左右されず、手法の強さを比較する。"],
            ["ATR", "平均的な値動きの大きさ。ATR14は直近14本の平均的な変動幅。", "急落幅、SL幅、高ボラ判定の基準に使う。"],
            ["PF", "総利益 ÷ 総損失。1.0超で利益側、1.3以上なら比較的良好。", "勝率だけでなく損益バランスを見る。"],
            ["DD", "ピークからどれだけ資産曲線が落ちたか。", "耐えられるリスクかを判断する。"],
            ["lookback", "過去何本を見るか。", "高値更新・安値更新の基準期間。"],
            ["exclude", "直近何本は水準タッチがないことを見る期間。", "直近で触られ続けている弱い水準を除外する。"],
        ],
        [1400, 4100, 3860],
    )

    add_heading(doc, "3. 主要結論", 1)
    add_bullet(doc, "TrendBreakV1 HYBRIDは10年検証で+191.53R、PF1.624。現時点の中核候補。")
    add_bullet(doc, "騙し回避では、ブレイク後3本以内に水準内へ戻るかどうかが強い特徴として出た。")
    add_bullet(doc, "同方向追加ポジションは、全体では最大2ポジションで+197.15R。大幅改善ではないが、XAUUSD/GBPJPY/CHFJPYでは効果あり。")
    add_bullet(doc, "Sai近似は全体ではH1が最も良好。H4/D1は一部手法だけ候補。")
    add_bullet(doc, "V字フィボはH4の61.8%戻しが強く、急落後V字の買い方向に絞るとPFが改善。")
    add_bullet(doc, "エリオット風5波は単体主力にはまだ弱く、補助仮説として扱うのが妥当。")

    add_heading(doc, "4. TrendBreakV1 ベースラインと騙し回避", 1)
    add_table(
        doc,
        ["項目", "結果"],
        [
            ["取引数", "461"],
            ["総R（コスト込み）", "+191.53R"],
            ["勝率", "36.88%"],
            ["PF", "1.624"],
            ["最大DD", "17.95R"],
            ["3本以内に終値がブレイク水準内へ戻った率", "37.53%"],
        ],
        [3600, 5760],
    )
    add_body(
        doc,
        "売買への活かし方: ブレイク後にすぐ水準内へ戻る動きは騙し候補です。"
        "単純に全て避けるのではなく、実体比率、終値位置、逆ヒゲ、停滞後ブレイクと組み合わせて使う方が自然です。",
    )
    add_table(
        doc,
        ["フィルタ", "意味", "取引数", "勝率", "総R", "PF", "最大DD"],
        [
            ["body_ratio>=0.6", "シグナル足の実体が強い", "377", "38.20%", "+176.27R", "1.716", "11.70R"],
            ["stagnation_then_break", "直前6本が停滞気味", "237", "39.66%", "+125.71R", "1.840", "14.64R"],
            ["confirm_3_closes_outside", "3本連続で水準外に残る", "287", "37.63%", "+129.06R", "1.684", "9.52R"],
        ],
        [2300, 2500, 900, 900, 1100, 800, 860],
    )

    add_heading(doc, "5. 同方向追加ポジション検証", 1)
    add_body(
        doc,
        "同方向に最大1から5ポジションまで持つ検証では、全体最適は最大2ポジションでした。"
        "ただし改善幅は小さく、無理に複雑化するほどではありません。",
    )
    add_table(
        doc,
        ["最大ポジション", "エントリー数", "勝率", "総R", "PF", "最大DD"],
        [
            ["1", "461", "36.88%", "+191.53R", "1.624", "18.27R"],
            ["2", "465", "37.20%", "+197.15R", "1.640", "18.27R"],
            ["3-5", "465", "37.20%", "+197.15R", "1.640", "18.27R"],
        ],
        [1600, 1600, 1200, 1600, 1200, 2160],
    )
    add_body(
        doc,
        "運用判断: XAUUSD、GBPJPY、CHFJPYは最大2ポジションを試す余地があります。"
        "USDJPY、EURJPY、SILVER、AUDJPYは基本1ポジション維持が無難です。",
    )

    add_heading(doc, "6. Sai近似ロジックの時間足比較", 1)
    add_table(
        doc,
        ["時間足", "取引数", "勝率", "総R", "PF", "最大連敗", "平均保有日数"],
        [
            ["H1", "3670", "41.63%", "+99.91R", "1.05", "24", "1.78"],
            ["H4", "1739", "39.91%", "-6.06R", "0.99", "16", "4.63"],
            ["D1", "292", "38.70%", "-11.42R", "0.94", "13", "17.78"],
        ],
        [1000, 1100, 1100, 1200, 900, 1200, 2860],
    )
    add_body(
        doc,
        "現時点では、Sai近似を丸ごとH4/D1へ移す根拠は弱いです。H1を軸にしつつ、"
        "H4のV字＋高値停滞、H4のレンジ抜け2回目以降、D1の値幅広め高値停滞だけを別枠で視覚確認するのが良さそうです。",
    )

    add_heading(doc, "7. V字フィボ・エリオット検証", 1)
    summary = pd.read_csv(ROOT / "backtests" / "elliott_fibo" / "results_2015_2024" / "summary_overall.csv")
    pick_names = [
        "VFIB_618_RR2",
        "VFIB_618_BODY50_RR2",
        "VFIB_618_BODY60_RR2",
        "VFIB_618_BODY60_LONG_RR2",
        "VFIB_618_BODY60_LONG_REC1_SPEED030_RR2",
        "ELLIOTT_W5_RR3_LOOSE",
    ]
    rows = []
    for name in pick_names:
        best = summary[summary["strategy"].eq(name)].sort_values("total_r_after_cost", ascending=False).head(1)
        if best.empty:
            continue
        row = best.iloc[0]
        rows.append(
            [
                str(row["timeframe"]),
                name,
                str(int(row["trades"])),
                fmt_pct(float(row["win_rate"])),
                fmt_r(float(row["total_r_after_cost"])),
                f"{float(row['pf_after_cost']):.2f}",
                f"{float(row['max_dd_r']):.2f}R",
            ]
        )
    add_table(
        doc,
        ["足", "戦略", "取引数", "勝率", "総R", "PF", "最大DD"],
        rows,
        [700, 3700, 850, 900, 1100, 800, 1310],
    )
    add_body(
        doc,
        "読み取り: V字フィボはH4の61.8%戻しが最も扱いやすい候補です。"
        "エリオット風5波はプラス例もありますがDDが重く、現段階では単体採用よりも環境認識の補助として扱う方が安全です。",
    )

    add_heading(doc, "8. 急落後V字回復の定量化", 1)
    add_body(
        doc,
        "急落後V字回復は、裁量で見つけると曖昧になりやすいため、次のように機械定義しました。",
    )
    add_bullet(doc, "確定スイング高値から確定スイング安値まで、H4なら3.5ATR以上落ちる。")
    add_bullet(doc, "その安値から下落幅の61.8%以上を終値で回復する。")
    add_bullet(doc, "シグナル足の実体がローソク足全体の60%以上。")
    add_bullet(doc, "買いのみ。SLはV字の谷より0.25ATR下、TPは2R。")
    add_bullet(doc, "補助指標として、急落速度、回復速度、下落幅ATRを記録する。")
    add_table(
        doc,
        ["条件", "足", "取引数", "勝率", "総R", "PF", "最大DD"],
        [
            ["H4 実体60%・急落後V字買い", "H4", "566", "43.29%", "+103.78R", "1.33", "14.62R"],
            ["H4 回復が下落本数以内", "H4", "402", "41.54%", "+61.86R", "1.27", "11.10R"],
            ["H4 急落幅4ATR・速度0.60以上", "H4", "187", "44.39%", "+41.52R", "1.41", "9.84R"],
            ["D1 実体60%・回復速度・急落速度", "D1", "103", "56.31%", "+34.14R", "1.82", "9.15R"],
        ],
        [3000, 700, 850, 900, 1100, 800, 2010],
    )
    add_body(
        doc,
        "運用候補: H4は取引数と総利益のバランスが良く、D1は回数は少ないものの勝率とPFが高いです。"
        "まずはH4の急落後V字をPineで可視化し、谷、61.8%ライン、シグナル足を目視確認するのが次の作業です。",
    )

    add_heading(doc, "9. 現時点の優先順位", 1)
    add_table(
        doc,
        ["優先度", "テーマ", "判断", "次のアクション"],
        [
            ["A", "TrendBreakV1 HYBRID", "主力候補。R管理がしやすい。", "Pineの約定/決済表示を安定化し、実運用に近いアラートを整備。"],
            ["A", "H4 急落後V字回復", "新しい有望候補。買い方向に絞る価値あり。", "PineでV字・61.8%・SL/TPを可視化。"],
            ["B", "騙し回避", "3本以内戻りと実体比率は有効な観察点。", "TrendBreakのフィルタとして採用するか追加検証。"],
            ["B", "同方向追加", "最大2ポジションは一部通貨で有効。", "通貨別にON/OFFを分ける。"],
            ["C", "エリオット風5波", "単体採用はまだ弱い。", "補助認識として使えるか検証継続。"],
        ],
        [900, 1800, 3100, 3560],
    )

    add_heading(doc, "10. 次にやること", 1)
    add_bullet(doc, "H4急落後V字回復のPine可視化スクリプトを作る。")
    add_bullet(doc, "TrendBreakV1に騙し回避フィルタを正式に組み込む前後で再検証する。")
    add_bullet(doc, "M5/M15は十分な過去データが入手できた後に同じ検証フレームで比較する。")
    add_bullet(doc, "各手法をGoogleスプレッドシートまたはCSVで、通貨別・年別・月別に管理できる形へ整理する。")

    add_heading(doc, "11. 参照ファイル", 1)
    add_body(doc, "主な生成物は以下のローカルパスに保存されています。")
    add_bullet(doc, str(ROOT / "backtests" / "trendbreak_v1" / "fakeout_feature_study_fast_2015_2024" / "report_ja.md"))
    add_bullet(doc, str(ROOT / "backtests" / "trendbreak_v1" / "pyramiding_sweep_2015_2024" / "report_ja.md"))
    add_bullet(doc, str(ROOT / "backtests" / "sai_h1" / "timeframe_comparison_2015_2024" / "report_ja.md"))
    add_bullet(doc, str(ROOT / "backtests" / "elliott_fibo" / "results_2015_2024" / "report_ja.md"))
    add_bullet(doc, str(ROOT / "backtests" / "elliott_fibo" / "results_2015_2024" / "vshape_quant" / "report_ja.md"))

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    make_doc()
    print(OUT)
